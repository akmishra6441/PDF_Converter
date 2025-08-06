import streamlit as st
import fitz  # PyMuPDF
from docx import Document
import io
import pytesseract
from PIL import Image

def convert_pdf_to_word(pdf_file_object, status_placeholder):
    """
    Core logic to convert a PDF to a Word document.
    Handles native text, images, and uses OCR as a fallback for scanned pages.

    Args:
        pdf_file_object: An in-memory file-like object from Streamlit's uploader.
        status_placeholder: A Streamlit placeholder to display status updates to the UI.

    Returns:
        An in-memory BytesIO stream of the resulting Word document, or None on failure.
    """
    try:
        # Open the PDF from the in-memory object
        pdf_bytes = pdf_file_object.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        st.error(f"Error opening PDF: {e}")
        return None

    word_doc = Document()
    total_pages = len(doc)

    for page_num in range(total_pages):
        status_placeholder.info(f"Processing page {page_num + 1} of {total_pages}...")
        page = doc.load_page(page_num)

        # Heuristic to decide if OCR is needed
        text = page.get_text()
        if len(text.strip()) < 100:  # Threshold can be adjusted
            status_placeholder.info(f"Page {page_num + 1} seems scanned. Applying OCR...")
            try:
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = pytesseract.image_to_string(img)
                word_doc.add_paragraph(ocr_text)
            except Exception as ocr_error:
                st.warning(f"Could not perform OCR on page {page_num + 1}. Skipping text. Error: {ocr_error}")
                word_doc.add_paragraph("[OCR failed for this page]")
        else:
            # If it's a native page, extract text blocks to preserve layout
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: b[1])
            for b in blocks:
                word_doc.add_paragraph(b[4])

        # Image Extraction
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_stream = io.BytesIO(image_bytes)
                word_doc.add_picture(image_stream)
            except Exception as img_error:
                st.warning(f"Could not add image {img_index + 1} from page {page_num + 1}: {img_error}")

        if page_num < total_pages - 1:
            word_doc.add_page_break()

    # Save the Word document to an in-memory stream
    doc_stream = io.BytesIO()
    word_doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream
